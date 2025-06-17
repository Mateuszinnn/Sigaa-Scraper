from selenium import webdriver
from selenium.common.exceptions import TimeoutException, NoSuchElementException, ElementClickInterceptedException
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import Select, WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from selenium.webdriver.chrome.service import Service
from selenium.webdriver.chrome.options import Options
from webdriver_manager.chrome import ChromeDriverManager
from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from collections import defaultdict
import re, time
from datetime import datetime, timedelta
import traceback
import sys

# === CONFIGURAÇÃO DO DRIVER ===
def configurar_driver():
    print("Configurando o driver do Chrome...")
    try:
        options = Options()
        options.add_argument("--headless")
        options.add_argument("--no-sandbox")
        options.add_argument("--disable-dev-shm-usage")
        options.add_argument("--disable-gpu")
        service = Service(ChromeDriverManager().install())
        driver = webdriver.Chrome(service=service, options=options)
        driver.get("https://sigaa.unb.br/sigaa/public/turmas/listar.jsf?aba=p-ensino")
        return driver, WebDriverWait(driver, 10)
    except Exception as e:
        print(f"ERRO ao configurar o driver: {str(e)}")
        print(traceback.format_exc())
        raise

# === FECHAR O MODAL DE COOKIES ===
def fechar_modal_cookies(wait):
    print("Verificando modal de cookies...")
    try:
        cookie_modal = wait.until(EC.visibility_of_element_located((By.ID, "sigaa-cookie-consent")))
        cookie_modal.find_element(By.XPATH, ".//button[contains(text(), 'Ciente')]").click()
        print("Modal de cookies fechado com sucesso.")
    except Exception as e:
        print(f"Aviso: Modal de cookies não encontrado ou erro ao fechar: {str(e)}")

# === INTERAÇÕES INICIAIS ===
def selecionar_departamento_por_indice(wait, index):
    try:
        Select(wait.until(EC.presence_of_element_located((By.ID, "formTurma:inputNivel")))).select_by_index(2)
        Select(wait.until(EC.presence_of_element_located((By.ID, "formTurma:inputDepto")))).select_by_index(index)
        wait.until(EC.element_to_be_clickable((By.XPATH, "//input[@value='Buscar']"))).click()
        time.sleep(5)
    except (TimeoutException, NoSuchElementException, ElementClickInterceptedException, IndexError) as e:
        print(f"[ERRO] Falha ao selecionar departamento por índice {index}: {e}")
    except Exception as e:
        print(f"[ERRO INESPERADO] {e}")

def selecionar_departamento_por_nome(wait, nome):
    try:
        Select(wait.until(EC.presence_of_element_located((By.ID, "formTurma:inputNivel")))).select_by_index(2)
        select_depto = Select(wait.until(EC.presence_of_element_located((By.ID, "formTurma:inputDepto"))))
        for option in select_depto.options:
            if nome.lower() in option.text.lower():
                select_depto.select_by_visible_text(option.text)
                break
        else:
            raise ValueError(f"Departamento com nome '{nome}' não encontrado.")
        
        wait.until(EC.element_to_be_clickable((By.XPATH, "//input[@value='Buscar']"))).click()
        time.sleep(5)
    except (TimeoutException, NoSuchElementException, ElementClickInterceptedException) as e:
        print(f"[ERRO] Falha ao selecionar departamento por nome '{nome}': {e}")
    except ValueError as e:
        print(f"[AVISO] {e}")
    except Exception as e:
        print(f"[ERRO INESPERADO] {e}")

# === LÓGICA DE CÓDIGOS DE HORÁRIO SUPORTANDO MULTI-TURNOS ===
HORARIOS_BASE = {
    'M': {1: "08:00", 2: "09:00", 3: "10:00", 4: "11:00", 5: "12:00"},
    'T': {1: "13:00", 2: "14:00", 3: "15:00", 4: "16:00", 5: "17:00", 6: "18:00"},
    'N': {1: "19:00"}
}

def agrupar_consecutivos_numeros(nums):
    try:
        if not nums:
            return []

        grupos = []
        temp = [nums[0]]

        for n in nums[1:]:
            if n == temp[-1] + 1:
                temp.append(n)
            else:
                grupos.append(temp)
                temp = [n]

        grupos.append(temp)
        return grupos

    except (IndexError, TypeError) as e:
        print(f"[ERRO] Falha ao agrupar números consecutivos: {e}")
        return []
    except Exception as e:
        print(f"[ERRO INESPERADO] Erro desconhecido ao agrupar: {e}")
        return []

def converter_codigo(codigo):
    try:
        pattern = re.compile(r"(\d+)([MTN]+)(\d+)")
        m = pattern.match(codigo.strip())
        if not m:
            return None

        dias_str, turnos_str, blocos_str = m.groups()
        dias_semana = {'2': 'Segunda', '3': 'Terça', '4': 'Quarta', '5': 'Quinta', '6': 'Sexta', '7': 'Sábado'}
        dias = [dias_semana[d] for d in dias_str if d in dias_semana]

        flatten = []
        for t in turnos_str:
            if t in HORARIOS_BASE:
                for b in sorted(HORARIOS_BASE[t].keys()):
                    flatten.append((t, b))

        indices = [int(ch) - 1 for ch in blocos_str if 0 <= int(ch) - 1 < len(flatten)]

        if not indices:
            return dias, []  # Nenhum horário válido encontrado

        grupos = agrupar_consecutivos_numeros(indices)

        periodos = []
        for g in grupos:
            start_idx = g[0]
            dur_min = (len(g) - 1) * 60 + 50
            t0, b0 = flatten[start_idx]
            inicio = datetime.strptime(HORARIOS_BASE[t0][b0], "%H:%M")
            fim = inicio + timedelta(minutes=dur_min)
            periodos.append(f"{inicio.strftime('%Hh%M')}–{fim.strftime('%Hh%M')}")

        return dias, periodos

    except (KeyError, ValueError, IndexError) as e:
        print(f"[ERRO] Código de horário malformado ou inválido: '{codigo}'. Detalhes: {e}")
        return None
    except Exception as e:
        print(f"[ERRO INESPERADO] Falha ao processar código '{codigo}': {e}")
        return None

# === EXTRAÇÃO DE DADOS ===
def extrair_dados(driver, apenas_fcte=False):
    print("Extraindo dados da tabela...")
    try:
        table = driver.find_element(By.CLASS_NAME, "listagem")
        rows = table.find_elements(By.TAG_NAME, "tr")
        cron = defaultdict(lambda: defaultdict(list))

        codigo_disciplina = ""
        nome_disciplina = ""

        for row in rows[1:]:
            try:
                # Atualiza o código e nome da disciplina (linha de título)
                titulo = row.find_element(By.CLASS_NAME, "tituloDisciplina").text.strip()
                partes = titulo.split(" ", 1)
                if len(partes) == 2:
                    codigo_disciplina, nome_disciplina = partes
                continue
            except Exception:
                pass  # Não é linha de título

            if not any(cl in row.get_attribute("class") for cl in ['linhaPar', 'linhaImpar']):
                continue

            try:
                cells = row.find_elements(By.TAG_NAME, "td")
                if len(cells) < 8:
                    continue

                raw_sala = cells[7].text.strip().upper()
                if apenas_fcte and not raw_sala.startswith(('FCTE', 'FGA')):
                    continue

                turma = cells[0].text.strip()
                professor = re.sub(r'\s*\(\d+h\)', '', cells[2].text.strip())
                professor = professor.lower().title()
                cods = re.findall(r'\d+[MTN]+\d+', cells[3].text.strip())

                clean = re.sub(r'^(?:FCTE|FGA)\s*-\s*', '', raw_sala)
                clean = re.sub(r'\s*\([^)]*\)', '', clean).strip()
                salas = [clean] if clean == 'NIT/LDS' else [s.strip() for s in clean.split('/')]

                for cod in cods:
                    try:
                        conv = converter_codigo(cod)
                        if not conv:
                            continue
                        dias, periodos = conv
                        for i, dia in enumerate(dias):
                            sala = f"FCTE - {salas[min(i, len(salas)-1)]}"
                            for per in periodos:
                                try:
                                    inicio, fim = per.split('–')
                                    cron[sala][dia].append({
                                        'inicio': inicio,
                                        'fim': fim,
                                        'codigo': codigo_disciplina,
                                        'turma': turma,
                                        'disciplina': nome_disciplina,
                                        'docente': professor
                                    })
                                except Exception as e:
                                    print(f"[ERRO] Falha ao dividir horário: {per} - {e}")
                    except Exception as e:
                        print(f"[ERRO] Erro ao converter código de horário '{cod}': {e}")

            except Exception as e:
                print(f"[ERRO] Falha ao processar linha de turma: {e}")

        return cron

    except Exception as e:
        print(f"[ERRO CRÍTICO] Falha ao extrair dados da tabela: {e}")
        return defaultdict(lambda: defaultdict(list))

# === CRIAÇÃO DO DOCUMENTO DINÂMICO ===
HORARIOS_FIXOS = [
    ("08h00", "09h50"),
    ("10h00", "11h50"),
    ("12h00", "13h50"),
    ("14h00", "15h50"),
    ("16h00", "17h50"),
    ("18h00", "19h50")
]

# Função utilitária para converter horário em minutos
def horario_para_minutos(horario):
    return int(horario[:2]) * 60 + int(horario[3:])

# Mapeia dias para colunas na tabela
dia_para_coluna = {
    "2": 1,  # Segunda
    "3": 2,  # Terça
    "4": 3,  # Quarta
    "5": 4,  # Quinta
    "6": 5,  # Sexta
    "7": 6   # Sábado
}

# === FONTE ===
def set_font_times_new_roman(doc):
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    rFonts = style.element.rPr.rFonts
    rFonts.set(qn('w:eastAsia'), 'Times New Roman')

# === MARGEM ===
def definir_margens(doc, cm_valor):
    for section in doc.sections:
        section.top_margin = Cm(cm_valor)
        section.bottom_margin = Cm(cm_valor)
        section.left_margin = Cm(cm_valor)
        section.right_margin = Cm(cm_valor)

def gerar_docx(cronogramas, filename="Mapa_de_Salas.docx"):
    print(f"Gerando DOCX: {filename}")
    print(f"Total de salas a processar: {len(cronogramas)}")
    doc = Document()
    set_font_times_new_roman(doc)
    definir_margens(doc, 0.5)
    for i, (sala, horarios) in enumerate(sorted(cronogramas.items()), start=1):
        print(f"[{i}/{len(cronogramas)}] Processando sala: {sala}")
        doc.add_heading(f"Sala {sala}", level=1)
        dias_semana_base = ["Horário", "Segunda", "Terça", "Quarta", "Quinta", "Sexta"]
        inclui_sabado = "Sábado" in horarios
        if inclui_sabado:
            dias_semana = dias_semana_base + ["Sábado"]
        else:
            dias_semana = dias_semana_base

        # cria tabela com número de colunas baseado em dias_semana
        tabela = doc.add_table(rows=7, cols=len(dias_semana))
        tabela.alignment = WD_TABLE_ALIGNMENT.CENTER
        tabela.style = 'Table Grid'
        tabela.autofit = False
        doc.add_page_break()

        for row in tabela.rows:
            row.cells[0].width = Cm(2)

        # 1) Cabeçalho fixo com dias
        for col_idx, titulo in enumerate(dias_semana):
            cell = tabela.cell(0, col_idx)
            cell.text = titulo
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            p.runs[0].bold = True

        # 2) Primeira coluna: horários fixos, altura exata
        for i, (inicio, fim) in enumerate(HORARIOS_FIXOS, start=1):
            cell_h = tabela.cell(i, 0)
            p = cell_h.paragraphs[0]
            p.clear() 

            run_inicio = p.add_run(inicio)
            run_inicio.bold = True
            p.add_run('\n')

            run_entre = p.add_run("às")
            run_entre.bold = True
            p.add_run('\n')

            run_fim = p.add_run(fim)
            run_fim.bold = True

            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            cell_h.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        # 3) Preencher aulas
        for dia_nome, aulas in horarios.items():
            if dia_nome not in dias_semana:
                continue
            col = dias_semana.index(dia_nome)
            for aula in aulas:
                mi = horario_para_minutos(aula['inicio'])
                mf = horario_para_minutos(aula['fim'])
                for idx, (bi_str, bf_str) in enumerate(HORARIOS_FIXOS, start=1):
                    bi = horario_para_minutos(bi_str)
                    bf = horario_para_minutos(bf_str)
                    if mi < bf and mf > bi:
                        cell = tabela.cell(idx, col)
                        # calcula duração do segmento dentro desta célula
                        overlap_inicio = max(mi, bi)
                        overlap_fim    = min(mf, bf)
                        dur = overlap_fim - overlap_inicio  # em minutos

                        if dur >= (bf - bi):  
                            # ocupa a célula inteira
                            p = cell.paragraphs[0]
                            p.clear()

                            # Código da disciplina (negrito, 14pt)
                            run_codigo = p.add_run(aula['codigo'])
                            run_codigo.bold = True
                            run_codigo.font.size = Pt(14)

                            p.add_run('\n')

                            # TURMA em negrito, tamanho 13
                            run_turma = p.add_run(f"TURMA {aula['turma']}")
                            run_turma.bold = True
                            run_turma.font.size = Pt(13)

                            p.add_run('\n')

                            # Nome da disciplina em itálico, tamanho 12
                            run_disc = p.add_run(aula['disciplina'].lstrip('-').strip().lower().title())
                            run_disc.italic = True
                            run_disc.font.size = Pt(12)

                            p.add_run('\n')

                            # Nome do professor já como você faz:
                            run_prof = p.add_run(f"Prof. {aula['docente']}")
                            run_prof.font.size = Pt(10)

                            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

                        else:
                            p = cell.paragraphs[0]
                            p.clear()

                            run_codigo = p.add_run(aula['codigo'])
                            run_codigo.bold = True
                            run_codigo.font.size = Pt(14)

                            p.add_run('\n')

                            # TURMA em negrito, tamanho 13
                            run_turma = p.add_run(f"TURMA {aula['turma']}")
                            run_turma.bold = True
                            run_turma.font.size = Pt(13)

                            p.add_run('\n')

                            # Nome da disciplina em itálico, tamanho 12
                            run_disc = p.add_run(aula['disciplina'].lstrip('-').strip().lower().title())
                            run_disc.italic = True
                            run_disc.font.size = Pt(12)

                            p.add_run('\n')

                            run_prof = p.add_run(f"Prof. {aula['docente']}")
                            run_prof.font.size = Pt(10)

                            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        # Espaço para próxima sala
        doc.add_paragraph()

    doc.save(filename)
    print(f"Documento gerado: {filename}")

# === MAIN ===
import traceback

def executar_scraping():
    try:
        print("Iniciando processo de scraping do SIGAA...")
        driver, wait = configurar_driver()
        fechar_modal_cookies(wait)
        selecionar_departamento_por_indice(wait, 2)
        cron_main = extrair_dados(driver)
        
        departamentos = [
            "CAMPUS UNB GAMA: FACULDADE DE CIÊNCIAS E TECNOLOGIAS EM ENGENHARIA - BRASÍLIA",
            "INSTITUTO DE FÍSICA - BRASÍLIA",
            "INSTITUTO DE QUÍMICA - BRASÍLIA",
            "DEPARTAMENTO DE MATEMÁTICA - BRASÍLIA",
            "DEPARTAMENTO DE ENGENHARIA MECANICA - BRASÍLIA",
            "DEPTO CIÊNCIAS DA COMPUTAÇÃO - BRASÍLIA"
        ]
        
        for i, depto in enumerate(departamentos, 1):
            print(f"\n[{i}/{len(departamentos)}] Processando departamento: {depto}")
            selecionar_departamento_por_nome(wait, depto)
            cron_ex = extrair_dados(driver, apenas_fcte=True)
            for s, d in cron_ex.items():
                for dia, aulas in d.items():
                    cron_main[s][dia].extend(aulas)
        
        arquivo = "Mapa_de_Salas.docx"
        gerar_docx(cron_main, arquivo)
        driver.quit()
        print("Processo de scraping concluído com sucesso!")
        return True, arquivo
    
    except Exception as e:
        print(f"Erro durante o scraping: {str(e)}")
        return False, str(e)

# Execução direta do script
if __name__ == "__main__":
    try:
        success, result = executar_scraping()
        if success:
            sys.exit(0)
        else:
            print(f"Erro: {result}")
            sys.exit(1)
    except Exception as e:
        print(f"Erro crítico: {str(e)}")
        print(traceback.format_exc())
        sys.exit(1)