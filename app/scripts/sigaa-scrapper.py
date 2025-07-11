from selenium import webdriver
from selenium.common.exceptions import TimeoutException, NoSuchElementException, ElementClickInterceptedException
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import Select, WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from selenium.webdriver.chrome.service import Service
from selenium.webdriver.chrome.options import Options
from webdriver_manager.chrome import ChromeDriverManager
from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from collections import defaultdict
import re, time
from datetime import datetime, timedelta
import traceback
import sys
import io
import os
from bs4 import BeautifulSoup

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
sys.stderr = io.TextIOWrapper(sys.stderr.buffer, encoding='utf-8')
sys.stdin.reconfigure(encoding='utf-8')
sys.stdout.reconfigure(encoding='utf-8', line_buffering=True)
sys.stderr.reconfigure(encoding='utf-8', line_buffering=True)
sys.stdout = os.fdopen(sys.stdout.fileno(), 'w', buffering=1)
sys.stderr = os.fdopen(sys.stderr.fileno(), 'w', buffering=1)

# === CONFIGURAÇÃO DO DRIVER ===
def configurar_driver():
    print("Configurando o driver do Chrome...")
    try:
        options = Options()
        options.add_argument("--headless")
        options.add_argument("--no-sandbox")
        options.add_argument("--disable-dev-shm-usage")
        options.add_argument("--disable-gpu")
        service = Service(ChromeDriverManager().install())
        driver = webdriver.Chrome(service=service, options=options)
        driver.get("https://sigaa.unb.br/sigaa/public/turmas/listar.jsf?aba=p-ensino")
        return driver, WebDriverWait(driver, 10)
    except Exception as e:
        print(f"ERRO ao configurar o driver: {str(e)}")
        print(traceback.format_exc())
        raise

# === FECHAR O MODAL DE COOKIES ===
def fechar_modal_cookies(wait):
    print("Verificando modal de cookies...")
    try:
        cookie_modal = wait.until(EC.visibility_of_element_located((By.ID, "sigaa-cookie-consent")))
        cookie_modal.find_element(By.XPATH, ".//button[contains(text(), 'Ciente')]").click()
        print("Modal de cookies fechado com sucesso.")
    except Exception as e:
        print(f"Aviso: Modal de cookies não encontrado ou erro ao fechar: {str(e)}")

# === INTERAÇÕES INICIAIS ===
def selecionar_departamento_por_indice(wait, index):
    print("[1/6] Processando departamento:  CAMPUS UNB GAMA: FACULDADE DE CIÊNCIAS E TECNOLOGIAS EM ENGENHARIA - BRASÍLIA")
    try:
        Select(wait.until(EC.presence_of_element_located((By.ID, "formTurma:inputNivel")))).select_by_index(2)
        Select(wait.until(EC.presence_of_element_located((By.ID, "formTurma:inputDepto")))).select_by_index(index)
        wait.until(EC.element_to_be_clickable((By.XPATH, "//input[@value='Buscar']"))).click()
        time.sleep(1)
    except (TimeoutException, NoSuchElementException, ElementClickInterceptedException, IndexError) as e:
        print(f"[ERRO] Falha ao selecionar departamento por índice {index}: {e}")
    except Exception as e:
        print(f"[ERRO INESPERADO] {e}")

def selecionar_departamento_por_nome(wait, nome):
    try:
        Select(wait.until(EC.presence_of_element_located((By.ID, "formTurma:inputNivel")))).select_by_index(2)
        select_depto = Select(wait.until(EC.presence_of_element_located((By.ID, "formTurma:inputDepto"))))
        for option in select_depto.options:
            if nome.lower() in option.text.lower():
                select_depto.select_by_visible_text(option.text)
                break
        else:
            raise ValueError(f"Departamento com nome '{nome}' não encontrado.")
        
        wait.until(EC.element_to_be_clickable((By.XPATH, "//input[@value='Buscar']"))).click()
        time.sleep(1)
    except (TimeoutException, NoSuchElementException, ElementClickInterceptedException) as e:
        print(f"[ERRO] Falha ao selecionar departamento por nome '{nome}': {e}")
    except ValueError as e:
        print(f"[ERRO] {e}")
    except Exception as e:
        print(f"[ERRO INESPERADO] {e}")

def definir_ano_e_periodo(driver, wait, ano: str, periodo: str):
    try:
        print(f"Definindo parâmetros: Ano {ano}, Período {periodo}...")

        ano_input = wait.until(EC.presence_of_element_located((By.ID, "formTurma:inputAno")))
        ano_input.clear()
        ano_input.send_keys(ano)

        periodo_select = Select(wait.until(EC.presence_of_element_located((By.ID, "formTurma:inputPeriodo"))))
        periodo_select.select_by_value(periodo)

        print("Ano e período configurados corretamente.")
    except Exception as e:
        print(f"[ERRO] Falha ao definir ano/período: {e}")

# === LÓGICA DE CÓDIGOS DE HORÁRIO SUPORTANDO MULTI-TURNOS ===
HORARIOS_BASE = {
    'M': {1: "08:00", 2: "09:00", 3: "10:00", 4: "11:00", 5: "12:00"},
    'T': {1: "13:00", 2: "14:00", 3: "15:00", 4: "16:00", 5: "17:00", 6: "18:00"},
    'N': {1: "19:00"}
}

def agrupar_consecutivos_numeros(nums):
    try:
        if not nums:
            return []

        grupos = []
        temp = [nums[0]]

        for n in nums[1:]:
            if n == temp[-1] + 1:
                temp.append(n)
            else:
                grupos.append(temp)
                temp = [n]

        grupos.append(temp)
        return grupos

    except (IndexError, TypeError) as e:
        print(f"[ERRO] Falha ao agrupar números consecutivos: {e}")
        return []
    except Exception as e:
        print(f"[ERRO INESPERADO] Erro desconhecido ao agrupar: {e}")
        return []

def converter_codigo(codigo):
    try:
        pattern = re.compile(r"(\d+)([MTN]+)(\d+)")
        m = pattern.match(codigo.strip())
        if not m:
            return None

        dias_str, turnos_str, blocos_str = m.groups()
        dias_semana = {'2': 'Segunda', '3': 'Terça', '4': 'Quarta', '5': 'Quinta', '6': 'Sexta', '7': 'Sábado'}
        dias = [dias_semana[d] for d in dias_str if d in dias_semana]

        flatten = []
        for t in turnos_str:
            if t in HORARIOS_BASE:
                for b in sorted(HORARIOS_BASE[t].keys()):
                    flatten.append((t, b))

        indices = [int(ch) - 1 for ch in blocos_str if 0 <= int(ch) - 1 < len(flatten)]

        if not indices:
            return dias, []

        grupos = agrupar_consecutivos_numeros(indices)

        periodos = []
        for g in grupos:
            start_idx = g[0]
            dur_min = (len(g) - 1) * 60 + 50
            t0, b0 = flatten[start_idx]
            inicio = datetime.strptime(HORARIOS_BASE[t0][b0], "%H:%M")
            fim = inicio + timedelta(minutes=dur_min)
            periodos.append(f"{inicio.strftime('%Hh%M')}–{fim.strftime('%Hh%M')}")

        return dias, periodos

    except (KeyError, ValueError, IndexError) as e:
        print(f"[ERRO] Código de horário malformado ou inválido: '{codigo}'. Detalhes: {e}")
        return None
    except Exception as e:
        print(f"[ERRO INESPERADO] Falha ao processar código '{codigo}': {e}")
        return None

# === EXTRAÇÃO DE DADOS ===
def extrair_dados(driver, apenas_fcte=False):
    print("Extraindo dados da tabela...")
    try:
        table = driver.find_element(By.CLASS_NAME, "listagem")
        rows = table.find_elements(By.TAG_NAME, "tr")
        cron = defaultdict(lambda: defaultdict(list))

        codigo_disciplina = ""
        nome_disciplina = ""
        DIAS_ORDEM = ['Segunda', 'Terça', 'Quarta', 'Quinta', 'Sexta', 'Sábado']

        for row in rows[1:]:
            try:
                titulo = row.find_element(By.CLASS_NAME, "tituloDisciplina").text.strip()
                partes = titulo.split(" ", 1)
                if len(partes) == 2:
                    codigo_disciplina, nome_disciplina = partes
                continue
            except Exception:
                pass

            if not any(cl in row.get_attribute("class") for cl in ['linhaPar', 'linhaImpar']):
                continue

            try:
                cells = row.find_elements(By.TAG_NAME, "td")
                if len(cells) < 8:
                    continue

                raw_sala = cells[7].text.strip().upper()
                if apenas_fcte and not raw_sala.startswith(('FCTE', 'FGA')):
                    continue

                turma = cells[0].text.strip()
                html_professores = cells[2].get_attribute("innerHTML")
                soup = BeautifulSoup(html_professores, "html.parser")
                linhas = soup.stripped_strings
                lista_professores = [
                    re.sub(r'\s*\(\d+h\)', '', linha).strip().lower().title()
                    for linha in linhas if linha.strip()
                ]
                professores_str = ", ".join(lista_professores)
                cods = re.findall(r'\d+[MTN]+\d+', cells[3].text.strip())

                clean = re.sub(r'^(?:FCTE|FGA)\s*-\s*', '', raw_sala)
                clean = re.sub(r'\s*\([^)]*\)', '', clean).strip()
                salas = [clean] if clean == 'NIT/LDS' else [s.strip() for s in clean.split('/')]             
                   
                eventos = []
                for cod_idx, cod in enumerate(cods):
                    convertido = converter_codigo(cod)
                    if not convertido:
                        continue
                    dias_list, periodos_list = convertido
                    for dia in dias_list:
                        for per in periodos_list:
                            inicio_str = per.split('–')[0].strip()
                            try:
                                if 'h' in inicio_str:
                                    h, m = inicio_str.split('h')
                                    total_min = int(h)*60 + int(m)
                                else:
                                    total_min = int(inicio_str)*60
                            except:
                                total_min = 0
                            eventos.append((dia, per, total_min, cod_idx))

                eventos_ordenados = sorted(
                    eventos, 
                    key=lambda x: (DIAS_ORDEM.index(x[0]) if x[0] in DIAS_ORDEM else 99, x[2])
                )
                
                num_eventos = len(eventos_ordenados)
                dias_unicos = sorted(set(e[0] for e in eventos_ordenados), 
                                    key=lambda d: DIAS_ORDEM.index(d))
                num_dias = len(dias_unicos)

                # Verificar se a disciplina é inconsistente
                if not (
                    len(salas) == len(cods) or
                    len(salas) == num_dias or
                    len(salas) == num_eventos or
                    len(salas) == 1
                ):
                    print(f"[DISCIPLINA_INCONSISTENTE] Código: {codigo_disciplina}, Turma: {turma}, Nome: {nome_disciplina}, Docente: {professores_str}, Códigos de horário: {cods}, Salas detectadas: {salas}, Dias únicos: {dias_unicos}")
                    print("[DISCIPLINA_INCONSISTENTE] Nenhuma das combinações esperadas é válida, significa que o número de salas não corresponde à quantidade de códigos de horários, dias ou eventos da disciplina, e também não é uma única sala.")
                    print("[DISCIPLINA_INCONSISTENTE] Favor verificar turma e adicionar ao documento manualmente!")
                    continue

                for idx, evento in enumerate(eventos_ordenados):
                    dia, per, _, cod_idx = evento

                    if len(salas) == len(cods):
                        sala = salas[cod_idx]
                    elif len(salas) == num_dias:
                        sala_idx = dias_unicos.index(dia)
                        sala = salas[sala_idx]
                    elif len(salas) == num_eventos:
                        sala = salas[idx]
                    elif len(salas) == 1:
                        sala = salas[0]

                    inicio, fim = per.split('–')
                    sala_completa = f"FCTE - {sala.strip()}"
                    evento_novo = {
                        'inicio': inicio.strip(),
                        'fim': fim.strip(),
                        'codigo': codigo_disciplina,
                        'turma': turma,
                        'disciplina': nome_disciplina,
                        'docente': professores_str
                    }

                    # Verifica se existe evento anterior contínuo e idêntico (exceto horários)
                    eventos_dia = cron[sala_completa][dia]
                    if eventos_dia:
                        ultimo = eventos_dia[-1]
                        if (
                            ultimo['codigo'] == evento_novo['codigo'] and
                            ultimo['turma'] == evento_novo['turma'] and
                            ultimo['disciplina'] == evento_novo['disciplina'] and
                            ultimo['docente'] == evento_novo['docente'] and
                            horario_para_minutos(ultimo['fim']) in [horario_para_minutos(evento_novo['inicio']) - 10,
                                                                    horario_para_minutos(evento_novo['inicio'])]
                        ):
                            # Mesclar horários: estende o fim
                            ultimo['fim'] = evento_novo['fim']
                        else:
                            eventos_dia.append(evento_novo)
                    else:
                        eventos_dia.append(evento_novo)
         
            except Exception as e:
                print(f"[ERRO] Falha ao processar linha de turma: {e}")
        return cron
    except NoSuchElementException:
        print("Não foram encontrados resultados para a busca com estes parâmetros.")
        return defaultdict(lambda: defaultdict(list))
    except Exception as e:
        print(f"[ERRO CRÍTICO] Falha ao extrair dados da tabela: {e}")
        return defaultdict(lambda: defaultdict(list))

# === CRIAÇÃO DO DOCUMENTO DINÂMICO ===
HORARIOS_FIXOS = [
    ("08h00", "09h50"),
    ("10h00", "11h50"),
    ("12h00", "13h50"),
    ("14h00", "15h50"),
    ("16h00", "17h50"),
    ("18h00", "19h50")
]

# Função utilitária para converter horário em minutos
def horario_para_minutos(horario):
    return int(horario[:2]) * 60 + int(horario[3:])

# === FONTE ===
def set_font_times_new_roman(doc):
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    rFonts = style.element.rPr.rFonts
    rFonts.set(qn('w:eastAsia'), 'Times New Roman')

# === MARGEM ===
def definir_margens(doc, cm_valor):
    for section in doc.sections:
        section.top_margin = Cm(cm_valor)
        section.bottom_margin = Cm(cm_valor)
        section.left_margin = Cm(cm_valor)
        section.right_margin = Cm(cm_valor)

ABBR_DISCIPLINAS = {
    "Engenharia": "Eng.",
    "Eletrônica": "Eletr.",
    "Eletromagnetismo": "Elet.",
    "Aplicada": "Apl.",
    "Especiais": "Esp.",
    "Tópicos": "Tóp.",
    "Projeto": "Proj.",
    "Computadores": "Comp.",
    "Programação": "Prog.",
    "Sistemas": "Sist.",
    "Teoria": "Teor.",
    "Fundamentos": "Fund.",
    "Métodos": "Met.",
    "Método": "Metó.",
    "Estruturas": "Estrut.",
    "Aeroespacial": "Aero.",
    "Dispositivos": "Disp.",
    "Energia": "Energ.",
    "Digital": "Disc.",
    "Desenvolvimento": "Desenv.",
    "Desenho": "Des.",
    "para": "p/",
    "Laboratório" : "Lab."
}

MAPEAMENTO_SALAS_COMPLETAS = {
    "AUDITÓRIO": "AUDITÓRIO – 245 – UAC",
    "CONTAINER 04": "CONTAINER Nº 04 – 30 capacidade ",
    "CONTEINER 04 - FUND VEÍCULOS ELÉTRICOS": "CONTAINER Nº 04 – 30 capacidade ",
    "I1": "SALA I-1 – AT-49/41 (TV) – 45 – UAC",
    "I2": "SALA I-2 – AT-42/48 – 70 – UAC",
    "I3": "SALA I-3 – AT-39/48 – 70 – UAC",
    "I4": "SALA I-4 – AT-32/41 (TV) – 45 – UAC",
    "I5": "SALA I-5 – AT-29/41 (TV) – 45 – UAC",
    "I6": "SALA I-6 – AT-22/48 LAB (Laptop) – 49 – UAC",
    "I7": "SALA I-7 – AT-19/48 LAB (Laptop/Desktop) – 48 – UAC",
    "I8": "SALA I-8 – AT-12/41 (TV) – 45 – UAC",
    "I9": "SALA I-9 – AT-09/41 – 130 – UAC",
    "I10": "SALA I-10 – AT-09/23 (LAB) – 80 – UAC",
    "LAB ELET": "LAB. ELETRICIDADE – 20 – UED",
    "LAB FÍSICA 1": "LAB. FÍSICA 1 – 25 - UED",
    "LAB MATERIAIS": "LAB. MATERIAIS – 15 – UED",
    "LAB NEI 1": "LAB. NEI 1 – 20 – UED",
    "LAB NEI 2": "LAB. NEI 2 – 20 – UED",
    "LAB OND": "LAB. FÍSICA 2 – 25 – UED",
    "LAB QUIMICA": "LAB. QUÍMICA – 20 – UED",
    "LAB SHP": "CONTAINER Nº 08 – LAB SHP",
    "LAB SS": "LAB. SS – 35 – UED",
    "LAB TERM": "LAB. TERMODINÂMICA – 25 – UED",
    "LAB TERMD.": "LAB. TERMODINÂMICA – 25 – UED",
    "LAB TERMOFLUIDOS": "LAB. TERMOFLUIDOS – 25 – UED",
    "LAB. FÍSICA 1": "LAB. FÍSICA 1 – 25 - UED",
    "LAB. MOCAP": "LAB. MOCAP – 80 – UED",
    "LAB. TERMOD.": "LAB. TERMODINÂMICA – 25 – UED",
    "LABORATÓRIO DE ELETRICIDADE": "LAB. ELETRICIDADE – 20 – UED",
    "LDTEA 302": "LDTEA – Sala 302 – 25 capacidade",
    "LDTEA 303": "LDTEA – Sala 303 – 25 capacidade",
    "MOCAP": "LAB. MOCAP – 80 – UED",
    "MULTIUSO": "SALA MULTIUSO – 25 – UAC",
    "NIT/LDS": "LAB NIT/LDS – 12 – UED",
    "S1": "SALA S-1 – A1-62/41 – 130 – UAC",
    "S2": "SALA S-2 – A1-59/41 – 130 – UAC",
    "S3": "SALA S-3 – A1-42/41 – 130 – UAC",
    "S4": "SALA S-4 – A1-32/41 – 130 – UAC",
    "S5": "SALA S-5 – A1-29/41 (TV) – 45 – UAC",
    "S6": "SALA S-6 – A1-22/48 – 70 – UAC",
    "S7": "SALA S-7 – A1-19/48 – 70 – UAC",
    "S8": "SALA S-8 – A1-12/41 (TV) – 45 – UAC",
    "S9": "SALA S-9 – A1-09/41 – 130 – UAC",
    "S10": "SALA S-10 – A1-09/23 (LAB) – 80 – UAC",
    "TERMO SUP": "LAB. TERMOFLUIDOS – 25 – UED"
}

def abbreviar_disciplina(nome_completo: str) -> str:
    """
    Substitui ocorrências de termos completos por suas abreviações (case-insensitive).
    """
    for completo, abbr in ABBR_DISCIPLINAS.items():
        pattern = rf"(?i)\b{re.escape(completo)}\b"
        nome_completo = re.sub(pattern, abbr, nome_completo)
    return nome_completo

def primeiro_ultimo_nome(nome_completo):
    partes = nome_completo.strip().split()
    if not partes:
        return ""
    if len(partes) == 1:
        return partes[0]
    return f"{partes[0]} {partes[-1]}"

def nome_sala_completo(sala: str) -> str:
    # Prioriza igualdade exata
    if sala in MAPEAMENTO_SALAS_COMPLETAS:
        return MAPEAMENTO_SALAS_COMPLETAS[sala]

    # Depois tenta correspondência parcial, priorizando chaves maiores
    for k in sorted(MAPEAMENTO_SALAS_COMPLETAS.keys(), key=len, reverse=True):
        if k in sala:
            return MAPEAMENTO_SALAS_COMPLETAS[k]
    return sala  # fallback

# === GERAÇÃO DO DOCX ===
def gerar_docx(cronogramas, filename="Mapa_de_Salas.docx"):
    print(f"Gerando DOCX: {filename}")
    print(f"Total de salas a processar: {len(cronogramas)}")
    doc = Document()
    section = doc.sections[0]
    header = section.header
    paragraph = header.paragraphs[0]
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    BASE_DIR = os.path.abspath(os.path.join(os.path.dirname(__file__), "..", ".."))
    imagem_path = os.path.join(BASE_DIR, "public", "Imagem1.png")

    run = paragraph.add_run()
    run.add_picture(imagem_path, width=Inches(6))

    set_font_times_new_roman(doc)
    definir_margens(doc, 0.5)

    for i, (sala, horarios) in enumerate(sorted(cronogramas.items()), start=1):
        print(f"[{i}/{len(cronogramas)}] Processando sala: {sala}")
        nome = nome_sala_completo(sala)
        par = doc.add_paragraph()
        par.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = par.add_run(nome)
        run.bold = True
        run.font.name = 'Times New Roman'
        run.font.size = Pt(14)

        dias_base = ["Horário", "Segunda", "Terça", "Quarta", "Quinta", "Sexta"]
        dias_semana = dias_base + (["Sábado"] if "Sábado" in horarios else [])

        tabela = doc.add_table(rows=1 + len(HORARIOS_FIXOS) * 2, cols=len(dias_semana))
        tabela.alignment = WD_TABLE_ALIGNMENT.CENTER
        tabela.style = 'Table Grid'
        tabela.autofit = False
        doc.add_page_break()

        for row in tabela.rows:
            row.cells[0].width = Cm(2.2)

        for idx, dia in enumerate(dias_semana):
            cell = tabela.cell(0, idx)
            cell.text = ""
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            run = p.add_run(dia.upper())
            run.bold = True
            run.font.size = Pt(10)

        for idx, (inicio, fim) in enumerate(HORARIOS_FIXOS):
            row1 = idx * 2 + 1
            row2 = idx * 2 + 2
            cell_h = tabela.cell(row1, 0).merge(tabela.cell(row2, 0))
            p = cell_h.paragraphs[0]
            p.clear()
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            run_inicio = p.add_run(inicio)
            run_inicio.bold = True
            run_inicio.font.size = Pt(16)
            p.add_run("\n")
            run_entre = p.add_run("às")
            run_entre.bold = True
            run_entre.font.size = Pt(16)
            p.add_run("\n")
            run_fim = p.add_run(fim)
            run_fim.bold = True
            run_fim.font.size = Pt(16)
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            cell_h.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        celulas_ocupadas = set()

        for dia, aulas in horarios.items():
            if dia not in dias_semana:
                continue
            col = dias_semana.index(dia)
            for aula in aulas:
                mi, mf = horario_para_minutos(aula['inicio']), horario_para_minutos(aula['fim'])
                for idx, (bi_str, bf_str) in enumerate(HORARIOS_FIXOS):
                    bi, bf = horario_para_minutos(bi_str), horario_para_minutos(bf_str)
                    blocos = [(bi, bi + 50), (bi + 50, bf)]
                    row1 = idx * 2 + 1
                    row2 = idx * 2 + 2

                    ocup1 = mi < blocos[0][1] and mf > blocos[0][0]
                    ocup2 = mi < blocos[1][1] and mf > blocos[1][0]

                    excecao_m5_t1 = mi == 720 and mf == 830 and bi == 720 
                    excecao_t6_n1 = mi == 1080 and mf == 1190 and bi == 1080  

                    if (ocup1 and ocup2) or excecao_m5_t1 or excecao_t6_n1:
                        cell = tabela.cell(row1, col).merge(tabela.cell(row2, col))
                        celulas_ocupadas.add((row1, col))
                        celulas_ocupadas.add((row2, col))
                    elif ocup1:
                        cell = tabela.cell(row1, col)
                        celulas_ocupadas.add((row1, col))
                    elif ocup2:
                        cell = tabela.cell(row2, col)
                        celulas_ocupadas.add((row2, col))
                    else:
                        continue

                    p = cell.paragraphs[0]
                    p.clear()
                    run_codigo = p.add_run(aula['codigo'])
                    run_codigo.bold = True
                    run_codigo.font.size = Pt(14)
                    p.add_run('\n')
                    run_turma = p.add_run(f"TURMA {aula['turma']}")
                    run_turma.bold = True
                    run_turma.font.size = Pt(13)
                    p.add_run('\n')
                    nome_original = aula['disciplina'].lstrip('-').strip().lower().title()
                    nome_abbr = abbreviar_disciplina(nome_original)
                    run_disc = p.add_run(nome_abbr)
                    run_disc.italic = True
                    run_disc.font.size = Pt(12)
                    p.add_run('\n')
                    nome_prof = primeiro_ultimo_nome(aula['docente'])
                    run_prof = p.add_run(f"Prof. {nome_prof}")
                    run_prof.font.size = Pt(10)
                    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        for idx in range(len(HORARIOS_FIXOS)):
            row1 = idx * 2 + 1
            row2 = idx * 2 + 2
            for col in range(1, len(dias_semana)):  
                if ((row1, col) not in celulas_ocupadas) and ((row2, col) not in celulas_ocupadas):
                    try:
                        tabela.cell(row1, col).merge(tabela.cell(row2, col))
                    except:
                        pass 
                    
        doc.add_paragraph()

    doc.save(filename)
    print(f"Documento gerado: {filename}")

# === MAIN ===
def executar_scraping():
    try:
        print("Iniciando processo de scraping do SIGAA...")
        driver, wait = configurar_driver()
        fechar_modal_cookies(wait)
        if len(sys.argv) > 2:
            ano = sys.argv[1]
            periodo = sys.argv[2]
        else:
            ano = input("Digite o ano (ex: 2025): ").strip()
            periodo = input("Digite o período (ex: 1 ou 2): ").strip()

        definir_ano_e_periodo(driver, wait, ano, periodo)
        selecionar_departamento_por_indice(wait, 2)

        cron_main = extrair_dados(driver)
        
        departamentos = [
            "INSTITUTO DE FÍSICA - BRASÍLIA",
            "INSTITUTO DE QUÍMICA - BRASÍLIA",
            "DEPARTAMENTO DE MATEMÁTICA - BRASÍLIA",
            "DEPARTAMENTO DE ENGENHARIA MECANICA - BRASÍLIA",
            "DEPTO CIÊNCIAS DA COMPUTAÇÃO - BRASÍLIA"
        ]
        
        for i, depto in enumerate(departamentos, 1):
            print(f"\n[{i+1}/{len(departamentos)+1}] Processando departamento: {depto}")
            selecionar_departamento_por_nome(wait, depto)
            cron_ex = extrair_dados(driver, apenas_fcte=True)
            for s, d in cron_ex.items():
                for dia, aulas in d.items():
                    cron_main[s][dia].extend(aulas)
        BASE_DIR = os.path.abspath(os.path.join(os.path.dirname(__file__), "..", ".."))
        arquivo = os.path.join(BASE_DIR, "public", "Mapa_de_Salas.docx")
        gerar_docx(cron_main, arquivo)
        driver.quit()
        print("Processo de scraping concluído com sucesso!")
        return True, arquivo
    
    except Exception as e:
        print(f"Erro durante o scraping: {str(e)}")
        return False, str(e)

# Execução direta do script
if __name__ == "__main__":
    try:
        success, result = executar_scraping()
        if success:
            sys.exit(0)
        else:
            print(f"Erro: {result}")
            sys.exit(1)
    except Exception as e:
        print(f"Erro crítico: {str(e)}")
        print(traceback.format_exc())
        sys.exit(1)

# implementado: 
# lista de abreviaturas de disciplinas, 
# primeiro e ultimo nome do prof, 
# se o dox já tinha sido gerado há a possibilidade de baixa-lo sem gerar novamente, 
# nome das salas por extenso, 
# imagem no cabeçalho, 
# logs para salas inconsistentes (nao acrescentar no word),
# mesclar turmas
